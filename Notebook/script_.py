import os
import logging
import datetime
import json
from dotenv import load_dotenv
from fastapi.responses import JSONResponse

from google.adk import Agent
from google.adk.agents import SequentialAgent, LoopAgent, ParallelAgent, LlmAgent
from google.adk.tools.tool_context import ToolContext
from google.genai import types
from google.adk.tools import FunctionTool
from google.adk.agents import LlmAgent
from google.adk.runners import Runner
from google.adk.sessions import InMemorySessionService
from google.genai import types


from pydantic import BaseModel, Field
from typing import List, Optional

import requests

import google.cloud.logging
from google.cloud import bigquery
from google.cloud import storage
from fastapi import FastAPI, Body

import io

# Setup logging
cloud_logging_client = google.cloud.logging.Client()
cloud_logging_client.setup_logging()
import uuid

load_dotenv()

model_name = os.getenv("MODEL")
print(model_name)

# --- 1. Define Constants ---
APP_NAME = "HR_SYSTEM_AGENT"
USER_ID = str(uuid.uuid4())
# "test_user_456"
SESSION_ID_TOOL_AGENT = "session_tool_agent_xyz"
# SESSION_ID_SCHEMA_AGENT = "session_schema_agent_xyz"
MODEL_NAME = "gemini-2.0-flash"


class File_Inputs(BaseModel):
    profile_path: str = Field(description="Resume file path to extract text from.")
    job_description: dict = Field(description="Job description to match against the resume in json format.")
# --- Utility Functions ---

def json_to_bigquery(json_file):
    client = bigquery.Client()
    dataset_id = 'candidate_cv'
    table_id = 'welders'
    table_ref = client.dataset(dataset_id).table(table_id)
    # Accept both dict and str input
    if isinstance(json_file, str):
        try:
            data = json.loads(json_file)
        except Exception as e:
            print(f"Error decoding JSON: {e}")
            return
    elif isinstance(json_file, dict):
        data = json_file
    else:
        print("Input to json_to_bigquery must be dict or JSON string.")
        return
    # Ensure data is a list of dictionaries
    if isinstance(data, dict):
        data = [data]
    elif not isinstance(data, list):
        print("Data for BigQuery must be a dict or list of dicts.")
        return
    errors = client.insert_rows_json(table_ref, data)
    if errors:
        print("Encountered errors while inserting rows: ", errors)
    else:
        print("Data successfully inserted into BigQuery.")

save_to_bq_tool = FunctionTool(func=json_to_bigquery)

def fetch_linkedin_profile(linkedin_url: str):
    """
    Fetch candidate info from LinkedIn using ProxyCurl API.
    """
    if not isinstance(linkedin_url, str) or not linkedin_url:
        return {"error": "Invalid LinkedIn URL."}
    # Uncomment and configure for real API usage
    # LINKEDIN_API_KEY = os.getenv("PROXYCURL_API_KEY")
    # if not LINKEDIN_API_KEY:
    #     return {"error": "LinkedIn API KEY not found."}
    # endpoint = 'https://nubela.co/proxycurl/api/v2/linkedin'
    # headers = {'Authorization': f'Bearer {LINKEDIN_API_KEY}'}
    # params = {'linkedin_profile_url': linkedin_url}
    # response = requests.get(endpoint, headers=headers, params=params)
    # if response.status_code == 200:
    #     return response.json()
    return {"error": "LinkedIn API KEY not found."}

fetch_linkedin_tool = FunctionTool(func=fetch_linkedin_profile)

def convert_gh_to_modified_json(input_json):
    if not isinstance(input_json, dict):
        raise ValueError("Input to convert_gh_to_modified_json must be a dict.")
    modified_json = {}
    profile = input_json.get("int_profile_data_json", {})
    if "skills" in profile and isinstance(profile["skills"], list):
        profile["skills"] = ", ".join([str(skill) for skill in profile["skills"]])
    modified_json.update(profile)
    attribute_scores = input_json.get("attribute_scores", {})
    modified_json.update(attribute_scores)
    semantic_match = input_json.get("semantic_match", {})
    modified_json.update(semantic_match)
    interview_data = input_json.get("interview_questions_json", {})
    if isinstance(interview_data, dict) and "interview_questions" in interview_data:
        questions = interview_data.get("interview_questions", [])
        for q in questions:
            if isinstance(q, dict):
                q.pop("rationale", None)
    modified_json.update(interview_data)
    final_ranking = input_json.get("final_ranking", {})
    for key, value in final_ranking.items():
        if key not in modified_json:
            modified_json[key] = value
        elif key == "ranking" and "ranking_summary" not in modified_json:
            modified_json["ranking_summary"] = final_ranking.get("summary", "")
    enriched_profile = input_json.get("enriched_profile_json", {})
    if "linkedin_profile" in enriched_profile:
        modified_json["linkedin_profile"] = enriched_profile["linkedin_profile"]
    flagged_gaps = input_json.get("flagged_gaps_json", {})
    flagged_issues = flagged_gaps.get("flagged_issues", [])
    if flagged_issues:
        modified_json["flagged_issue"] = ",".join([issue.get("issue", "") for issue in flagged_issues])
        modified_json["flagged_issue_description"] = ",".join([issue.get("description", "") for issue in flagged_issues])
        modified_json["flagged_issue_severity"] = ",".join([issue.get("severity", "") for issue in flagged_issues])
        modified_json["flagged_issue_resolution"] = ",".join([issue.get("resolution", "") for issue in flagged_issues])
    return modified_json

def save_context_to_json(context_data: dict, filename: str, output_key: str):
    base = os.path.splitext(os.path.basename(filename))[0]
    now_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    save_filename_json = f"{output_key}_{base}_{now_str}.json"
    if not isinstance(context_data, dict):
        try:
            context_data = json.loads(context_data)
        except Exception:
            context_data = {"data": str(context_data)}
    _new_context_data = convert_gh_to_modified_json(context_data)
    json_to_bigquery(_new_context_data)
    try:
        with open(save_filename_json, "w", encoding="utf-8") as f:
            json.dump(_new_context_data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        return f"Error saving context to {save_filename_json}: {e}"
    print(f"Saved context to {save_filename_json}")
    # print(type(_new_context_data))
    # print(json.dumps(_new_context_data))
    return _new_context_data


save_context_tool = FunctionTool(func=save_context_to_json)

def extract_text_from_file(file_path: str):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    if ext == ".pdf" or ext==".PDF":
        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            raise RuntimeError(f"Error reading PDF: {e}")
    elif ext == ".docx":
        try:
            from docx import Document
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            raise RuntimeError(f"Error reading DOCX: {e}")
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    return text


from google.cloud import storage
from PyPDF2 import PdfReader
import docx
import io

def extract_texts_from_gcs(bucket_name, prefix, service_account_json, archive_prefix="archive_cv/"):
    """
    Reads PDF and DOCX files from a GCP bucket, extracts text, and moves files to archive after processing.
    Returns a dict: {blob_name: extracted_text}
    """
    storage_client = storage.Client.from_service_account_json(service_account_json)
    gcs_bucket = storage_client.bucket(bucket_name)
    extracted = {}
    for blob in gcs_bucket.list_blobs(prefix=prefix):
        text = ""
        if blob.name.endswith('.pdf'):
            pdf_bytes = blob.download_as_bytes()
            pdf_stream = io.BytesIO(pdf_bytes)
            reader = PdfReader(pdf_stream)
            for page in reader.pages:
                text += page.extract_text() or ""
        elif blob.name.endswith('.docx'):
            try:
                file_bytes = blob.download_as_bytes()
                file_stream = io.BytesIO(file_bytes)
                doc = docx.Document(file_stream)
                text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                text = f"Error reading DOCX: {e}"
        if text:
            extracted[blob.name] = text
        # Move to archive after processing
        gcs_bucket.copy_blob(blob, gcs_bucket, f"{archive_prefix}{blob.name}")
        blob.delete()
    return extracted





from typing import List, Optional
from pydantic import BaseModel, Field

class EducationEntry(BaseModel):
    degree: str = Field(description="Degree or qualification obtained")
    institution: str = Field(description="Name of the educational institution")
    start_year: Optional[int] = Field(default=None, description="Year education started")
    end_year: Optional[int] = Field(default=None, description="Year education ended or expected to end")

class ExperienceEntry(BaseModel):
    job_title: str = Field(description="Job title or position held")
    company: str = Field(description="Name of the company or organization")
    start_date: Optional[str] = Field(default=None, description="Start date (YYYY-MM or similar)")
    end_date: Optional[str] = Field(default=None, description="End date or 'Present'")
    description: Optional[str] = Field(default=None, description="Brief description of responsibilities or achievements")

class CertificationEntry(BaseModel):
    name: str = Field(description="Name of the certification")
    provider: Optional[str] = Field(default=None, description="Name of the certification provider")
    certificate_url: Optional[str] = Field(default=None, description="URL to the certification")
    start_date: Optional[str] = Field(default=None, description="Start date of the certification (YYYY-MM or similar)")
    expire_date: Optional[str] = Field(default=None, description="Expiration date of the certification (YYYY-MM or similar)")


class ResumeOutput(BaseModel):
    name: str = Field(description="Full name of the candidate")
    email: Optional[str] = Field(default=None, description="Email address")
    phone: Optional[str] = Field(default=None, description="Phone number")
    address: Optional[str] = Field(default=None, description="Mailing address")
    location: Optional[str] = Field(default=None, description="Current location or city")
    summary: Optional[str] = Field(default=None, description="Professional summary or objective")
    skills: List[str] = Field(default_factory=list, description="List of skills")
    education: List[EducationEntry] = Field(default_factory=list, description="List of education entries")
    experience: List[ExperienceEntry] = Field(default_factory=list, description="List of work experience entries")
    certifications: List[CertificationEntry] = Field(default_factory=list, description="List of certifications")
    languages: List[str] = Field(default_factory=list, description="Languages spoken or known")
    projects: List[str] = Field(default_factory=list, description="Notable projects")
    linkedin: Optional[str] = Field(default=None, description="LinkedIn profile URL")
    github: Optional[str] = Field(default=None, description="GitHub profile URL")

extraction_agent = LlmAgent(
    name="file_path_data_extractor",
    # model=model_name,
    description="Receives a file path to a resume (PDF or DOCX), uses the extract_text_from_file tool to extract all readable text from the file, and returns the raw extracted text exactly as found in the document, without any modification, inference, or reformatting. Handles unsupported file types with an appropriate error message.",
    instruction="""
    INSTRUCTIONS:
    You will receive a user profile file path in the variable called profile_path.
    Your task is to extract all readable text from the provided local file path (PDF, DOC, or DOCX).
    - Use the 'extract_text_from_file' tool to process the file.
    - Return the extracted text exactly as found in the document, without any modification, inference, or reformatting.
    - If the file type is unsupported, return an appropriate error message.
    """,
    generate_content_config=types.GenerateContentConfig(
        temperature=0,
    ),
    input_schema=File_Inputs,
    tools=[FunctionTool(
            func=extract_text_from_file,
        )],
    output_key="extracted_text",
)


formatter_agent = LlmAgent(
    name="extraction_formatter",
    model=model_name,
    description="Converts unstructured resume text into structured JSON data using the ResumeOutput Pydantic schema.",
    instruction="""
You will receive resume text in the variable {{extracted_text}}.
Your task:
- Carefully parse the provided resume text and extract all relevant information *exactly* as it appears in the text. Do not add, infer, or modify any information.
- Format the extracted information as a JSON object that strictly matches the ResumeOutput Pydantic schema. Ensure every field in the schema is present in the output. Use empty strings, empty lists, or None where data is missing in the extracted text.
- If the input text does not appear to be a valid resume, or if you are unable to extract any information, return an appropriate error message.
- Output a JSON object with fields: 'int_profile_data_json' contains all the fields of output schema.
""",
    generate_content_config=types.GenerateContentConfig(
        temperature=0,
    ),
    output_schema=ResumeOutput,
    output_key="int_profile_data_json",
)


save_context_agent = LlmAgent(
    name="extraction_formatter",
    model=model_name,
    description=(
        "Saves the extracted profile data to a JSON file with the name 'Profile_Analysis_Result_.json'. "
        "The file will contain the structured resume data extracted from the provided text. "
    ),
    instruction="""
You will receive resume text in the variable {{merged_context_json}}.
Take filename from profile_path context variable and save the merged context JSON data to a file named 'Profile_Analysis_Result_<filename>_YYYYMMDD_HHMMSS.json'.
For saving pass it to tool 'save_context_to_json' with the following parameters:
    - context_data: the merged JSON object
    - filename: the original profile_path (or its_base name) if notable extract then use demo_.pdf
    - output_key: "Profile_Analysis_Result_json"
Ensure the file is saved in the current working directory.
and 
- The output from tool will be dictionary, just return the dictionary with as JSON
""",
    generate_content_config=types.GenerateContentConfig(
        temperature=0,
    ),
    tools=[save_context_tool],
    output_key="save_context",
)



merge_context_agent = LlmAgent(
    name="merge_context_agent",
    model=model_name,
    description=(
        "Merges all available context variables (JSON outputs from previous agents) into a single JSON object. "
        "Each merged entry is assigned a key corresponding to its context variable name. "
        "The following context variables will be merged: "
        "{{int_profile_data_json}}, {{attribute_scores}}, {{semantic_match}}, {{final_ranking}}, {{enriched_profile_json}}, {{flagged_gaps_json}}, {{interview_questions_json}}."
    ),
    instruction="""
You will receive multiple context variables, each containing JSON data from previous agents:
- {{int_profile_data_json}}
- {{attribute_scores}}
- {{semantic_match}}
- {{final_ranking}}
- {{enriched_profile_json}}
- {{flagged_gaps_json}}
- {{interview_questions_json}}

Your task:
- For each context variable, assign its variable name as the key and its JSON content as the value.
- Merge all such key-value pairs into a single JSON object.
- Return the merged JSON object.
- If any context variable is missing or empty, skip it.
""",
    generate_content_config=types.GenerateContentConfig(
        temperature=0,
    ),
    output_key="merged_context_json",

)


profile_formation_agent = SequentialAgent(
    name="resume_extraction_pipeline",
    description="Extracts text in sequential manner",
    sub_agents=[extraction_agent, formatter_agent],

)


jd_agent = LlmAgent(
    name="jd_parser",
    model=model_name,
    description="Parses the provided job description text and extracts structured requirements.",
    instruction="""
You will receive a job description in the variable job_description.
Your task:
- Carefully parse the job description and extract key requirements, skills, qualifications, and responsibilities *exactly* as they appear in the text. Do not add, infer, or modify any information.
- Format the extracted information as a JSON object.
- If the input is not a valid job description, or if you are unable to extract any information, return an appropriate error message.
""",
    generate_content_config=types.GenerateContentConfig(
        temperature=0,
    ),
    output_key="jd_json",
)

parallel_agent = ParallelAgent(
    name="resume_and_jd_parallel",
    description="Processes profile_path to profile_formation_agent and pass job_description key information to jd_agent in parallel.",
    sub_agents=[
        profile_formation_agent,
        jd_agent
    ],
)


match_agent = Agent(
    name="profile_jd_matcher",
    model=model_name,
    description="Compares structured resume data and job description to compute attribute match scores.",
    instruction="""
You will receive:
- Resume data in the variable {{int_profile_data_json}}
- Job description data in the variable {{jd_json}}

Your task:
- Compare the resume data and job description data.
- Compute and return a set of attribute scores (e.g., skill match, experience match, education match, overall fit).
- Output a JSON object with fields: 'attribute_scores' contains fields 'skill_match','experience_match','education_match','certification_match' and 'overall_fit' .

""",
    generate_content_config=types.GenerateContentConfig(
        temperature=0,
    ),
    output_key="attribute_scores",
    # tools = [save_context_tool]
)


semantic_scoring_agent = Agent(
    name="semantic_scoring_agent",
    model=model_name,
    description="Performs semantic scoring and ranking of candidate profiles against job requirements using job-role ontologies.",
    instruction="""
You will receive:
- Resume data in the variable {{int_profile_data_json}}
- Job description data in the variable {{jd_json}}

Your task:
- Analyze the candidate's qualifications and experience against the job requirements and project needs.
- Use semantic similarity and job-role ontologies to assess the match between the candidate and the job.
- Rank the candidate and provide a semantic match score (0-100), along with a brief explanation for the score.
- Output a JSON object with fields: 'semantic_score', 'ranking', and 'explanation'.
""",
    generate_content_config=types.GenerateContentConfig(
        temperature=0,
    ),
    output_key="semantic_match",
    # tools = [save_context_tool]
)


final_ranking_agent = Agent(
    name="final_ranking_agent",
    model=model_name,
    description="Aggregates attribute scores and semantic match to provide a final candidate ranking and summary.",
    instruction="""
You will receive:
- Attribute scores in the variable {{attribute_scores}}
- Semantic match in the variable {{semantic_match}}

Your task:
- Aggregate the attribute scores and semantic match score.
- Provide a final ranking, overall match percentage, and a summary explanation.
- Output a JSON object with fields: 'final_score', 'ranking', and 'summary'.
""",
    generate_content_config=types.GenerateContentConfig(
        temperature=0,
    ),
    output_key="final_ranking",
    # tools = [save_context_tool]
)


external_hr_enrichment_agent = Agent(
    name="external_hr_enrichment_agent",
    model=model_name,
    description="Checks for LinkedIn profile URL in the candidate profile and fetches additional info if present. If LinkedIn URL is not found, returns a single line indicating this.",
    instruction="""
You will receive structured resume data in the variable {{int_profile_data_json}}.
Your task:
- If a LinkedIn profile URL is present, use the 'fetch_linkedin_profile' tool to fetch additional candidate information and output the enriched profile as JSON.
- If a LinkedIn profile URL is NOT present, return a JSON object with a single field: {"linkedin_profile": "not found"}.
- Do not prompt or return any other information if the LinkedIn URL is missing.
""",
    generate_content_config=types.GenerateContentConfig(
        temperature=0,
    ),
    tools=[fetch_linkedin_tool],
    output_key="enriched_profile_json"
)


gap_flagging_agent = Agent(
    name="gap_flagging_agent",
    model=model_name,
    description="Flags mismatches or gaps in candidate profiles, such as expired certifications or missing skills, and displays the results in JSON format.",
    instruction="""
You will receive:
- Enriched candidate profile in the variable {{enriched_profile_json}}
- Job description data in the variable {{jd_json}}

Your task:
- Analyze the candidate profile for mismatches or gaps (e.g., expired certifications, missing or critical skills).
- Flag and list all issues found.
- Output a JSON object with a list of flagged issues and their details.
- Display the flagged gaps and issues in JSON format.
""",
    generate_content_config=types.GenerateContentConfig(
        temperature=0,
    ),
    output_key="flagged_gaps_json"
    )

interview_question_agent = Agent(
    name="interview_question_agent",
    model=model_name,
    description="Suggests interview questions or skill assessments based on the job role and candidate profile.",
    instruction="""
You will receive:
- Enriched candidate profile in the variable {{enriched_profile_json}}
- Job description data in the variable {{jd_json}}

Your task:
- Suggest relevant interview questions or skill assessments tailored to the job role and candidate's background.

- Output a JSON object with fields: 'interview_questions' and 'skill_assessments'.
- Each entry in the interview_questions array should be an object containing the subfields category and question.
- Similarly, each entry in the skill_assessments array should be an object containing the subfields assessment_type, description, and rationale.
""",
    generate_content_config=types.GenerateContentConfig(
        temperature=0,
    ),

    output_key="interview_questions_json"
)



pipeline_agent = SequentialAgent(
    name="advanced_profile_jd_scoring_pipeline",
    description="Runs the full candidate-job matching, enrichment, compliance, recommendation, and dashboard pipeline.",
    sub_agents=[
        parallel_agent,                  # Resume & JD extraction
        match_agent,                  # Attribute scoring
        semantic_scoring_agent,          # Semantic scoring
        final_ranking_agent,             # Final ranking
        external_hr_enrichment_agent,    # Enrich profile from HR systems
        gap_flagging_agent,              # Flag gaps/mismatches
        interview_question_agent,        # Suggest interview questions
        merge_context_agent,             # Merge all context variables into a single JSON object
        
        save_context_agent,              # Save the final merged context JSON
    ],


)





root_agent = Agent(
    name="root_agent",
    model=model_name,
    description=(
        "Coordinates the end-to-end candidate-job matching workflow: extracts and structures resume and job description, computes match scores, aggregates results, enriches candidate data, flags profile gaps, and displays all intermediate and final results."
    ),
    global_instruction=(
        "You are a candidate-job matching assistant. Accept a resume file path and job description, extract and structure both, compute and aggregate match scores, enrich the profile using external HR systems, flag gaps, and display all intermediate and final results."
    ),
    instruction="""
    1. Accept a resume file path and job description from the user.
    2. Store them as 'profile_path' and 'job_description'.
    3. Run the pipeline_agent to:
        a. Extract and structure resume and job description.
        a1. Display the extracted resume text and job description.
        b. Compute and display attribute and semantic match scores.
        c. Aggregate results for a final ranking.
        d. Enrich the profile using external HR systems.
        e. Flag and display profile gaps.
        f. Merge all context variables into a single JSON object.
        g. Save the final merged context JSON to a file named based on the profile_path.
    4. Return all intermediate and final results in JSON format.
    5. Display an error message if any step fails.

    6. Save the final merged context JSON to a file and Profile_path base  name as the filename.
    7. The final output will be a JSON object containing all the merged context variables, including:
       - Profile data
       - Attribute scores
       - Semantic match
       - Final ranking
       - Enriched profile data
       - Flagged gaps
       - Interview questions

    """,
    sub_agents=[pipeline_agent],
    output_key="merged_context_json",
)


# --- 5. Set up Session Management and Runners ---
session_service = InMemorySessionService()


agent_runner = Runner(
    agent=root_agent,
    app_name=APP_NAME,
    session_service=session_service
)


import re

def extract_json_from_code_block(text: str) -> str:
    """
    Extracts JSON string from a code block like ```json ... ```
    If no code block is found, returns the original text.
    """
    if not isinstance(text, str):
        return text
    # Match ```json ... ```
    # match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", text, re.IGNORECASE)
    text_ = text.replace("```json", "").replace("```", "")
    print(text_)
    # match = re.search(r"```json\s*([\s\S]*?)\s*)
    # if match:
    #     return match.group(1).strip()
    return text_


app = FastAPI(
    title="Agent Comparison API",
    description="API to compare LlmAgent with Tools vs. Output Schema",
    version="1.0.0"
)


@app.on_event("startup")
async def startup_event():
    # Always (re-)create sessions at startup
    await session_service.create_session(app_name=APP_NAME, user_id=USER_ID, session_id=SESSION_ID_TOOL_AGENT)
    # await session_service.create_session(app_name=APP_NAME, user_id=USER_ID, session_id=SESSION_ID_SCHEMA_AGENT)

@app.post("/multi_agent_call", summary="Get capital using tool agent")
async def get_capital_with_tool(request: File_Inputs):
    query_json = json.dumps(request.dict())
    result = {}

    async def run():
        user_content = types.Content(role='user', parts=[types.Part(text=query_json)])
        final_response_content = "No final response received."
        async for event in agent_runner.run_async(user_id=USER_ID, session_id=SESSION_ID_TOOL_AGENT, new_message=user_content):
            if event.is_final_response() and event.content and event.content.parts:
                final_response_content = event.content.parts[0].text
        return final_response_content

    response = await run()
    # Get session state
    current_session = await session_service.get_session(app_name=APP_NAME, user_id=USER_ID, session_id=SESSION_ID_TOOL_AGENT)
    stored_output = current_session.state.get(root_agent.output_key)
    try:
        print("IN Try")
        parsed_output = json.loads(stored_output)
        result["session_state"] = parsed_output
    except Exception:
        print("IN Except")
        result["session_state"] = stored_output
    print("OUT")
    response2 = response.replace("```json", "").replace("```", "")
    print(response2)
    safe_response2 = response2.replace('\\', '\\\\')
    response_json_str = json.dumps(json.loads(safe_response2), indent=4)
    # Replace single backslashes with double backslashes
    print(response_json_str)
    result["response"] = response2
    # import json
    response_1 = json.loads(response_json_str)
    # return result
    return JSONResponse(content={"response" : response_1["save_context_to_json_response"]})

